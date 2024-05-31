import streamlit as st
import nltk
from nltk import ne_chunk, pos_tag, word_tokenize
from nltk.sentiment import SentimentIntensityAnalyzer
from googletrans import Translator
from nltk.tree import Tree
from nltk.chunk import RegexpParser
import docx
import pyaudio
import wave
import pyttsx3
import os
import subprocess
import openai
import speech_recognition as sr

################################################################
# Programa de Grabar Audio y Convertirlo a Texto
def grabar_audio_texto():
    st.header("Grabar Audio y Convertir a Texto")
    duracion = st.number_input("Duración de la grabación (segundos):", min_value=1, max_value=60, value=5)

    if st.button("Grabar"):
        nombre_archivo = "grabacion.wav"
        grabar_audio(nombre_archivo, duracion)
        texto_reconocido = convertir_audio_a_texto(nombre_archivo)
        st.write("Texto reconocido:")
        st.write(texto_reconocido)

def grabar_audio(nombre_archivo, duracion):
    formato = pyaudio.paInt16
    canales = 1
    tasa_muestreo = 44100
    tamano_buffer = 1024

    audio = pyaudio.PyAudio()

    stream = audio.open(format=formato, channels=canales, rate=tasa_muestreo, input=True, frames_per_buffer=tamano_buffer)
    frames = []

    for _ in range(0, int(tasa_muestreo / tamano_buffer * duracion)):
        data = stream.read(tamano_buffer)
        frames.append(data)

    stream.stop_stream()
    stream.close()
    audio.terminate()

    archivo_salida = wave.open(nombre_archivo, 'wb')
    archivo_salida.setnchannels(canales)
    archivo_salida.setsampwidth(audio.get_sample_size(formato))
    archivo_salida.setframerate(tasa_muestreo)
    archivo_salida.writeframes(b''.join(frames))
    archivo_salida.close()

def convertir_audio_a_texto(nombre_archivo):
    reconocedor = sr.Recognizer()

    with sr.AudioFile(nombre_archivo) as fuente:
        audio = reconocedor.record(fuente)

    try:
        texto = reconocedor.recognize_google(audio, language='es-ES')
        return texto
    except sr.UnknownValueError:
        return "No se pudo reconocer el audio."
    except sr.RequestError as e:
        return "Error al solicitar resultados al servicio de reconocimiento de voz de Google: {0}".format(e)


################################################################
def texto_audio(): 
    def texto_a_voz(texto, nombre_archivo):
        engine = pyttsx3.init()
        engine.save_to_file(texto, nombre_archivo)
        engine.runAndWait()

    texto = st.text_area("Ingrese el texto para convertir a audio")
    nombre_archivo = st.text_input("Nombre del archivo de salida (sin extensión)")
    if st.button("Convertir a Audio"):
        if texto and nombre_archivo:
            texto_a_voz(texto, nombre_archivo + ".wav")
            st.write(f"El archivo de audio \"{nombre_archivo}.wav\" ha sido guardado correctamente.")
        else:
            st.write("Por favor, proporciona tanto el texto como el nombre del archivo.")

def audio_texto_audio():
    def texto_a_audio(texto):
        engine = pyttsx3.init()
        engine.setProperty('rate', 150)
        engine.save_to_file(texto, 'audio_generado3.mp3')
        engine.runAndWait()

    def leer_archivo_texto(archivo):
        extension = os.path.splitext(archivo)[1]
        if extension == '.txt':
            with open(archivo, 'r', encoding='utf-8') as f:
                texto = f.read()
            return texto
        elif extension == '.docx':
            doc = docx.Document(archivo)
            texto = ''
            for paragraph in doc.paragraphs:
                texto += paragraph.text + '\n'
            return texto
        else:
            st.write("Formato de archivo no compatible.")

    archivo = st.file_uploader("Sube tu archivo (.txt o .docx)", type=["txt", "docx"])
    if st.button("Convertir a Audio"):
        if archivo:
            texto = leer_archivo_texto(archivo)
            if texto:
                texto_a_audio(texto)
                st.write("Texto convertido a audio y reproducido.")
            else:
                st.write("No se pudo leer el archivo.")
        else:
            st.write("Por favor, sube un archivo.")

def chatGPT():
    openai.api_key = 'sk-proj-0vy7ddCyHr2D6xAHUNqPT3BlbkFJoJa85W1JpDP3v6sgJCkW'  # key
    def enviar_solicitud_chat(mensaje):
        respuesta = openai.Completion.create(
            engine='gpt-3.5-turbo-instruct',
            prompt=mensaje,
            max_tokens=50,
            temperature=0.7
        )
        return respuesta.choices[0].text.strip()

    mensaje = st.text_area("Escribe tu mensaje para ChatGPT")
    if st.button("Enviar"):
        if mensaje:
            respuesta = enviar_solicitud_chat(mensaje)
            st.write("Respuesta de ChatGPT: " + respuesta)
        else:
            st.write("Por favor, escribe un mensaje.")

def analizador_sentiment():
    ola = SentimentIntensityAnalyzer()
    translator = Translator()

    def traducir_to_eng(text):
        translated_text = translator.translate(text, dest='en').text
        return translated_text

    frase = st.text_input("Introduce una frase para analizar")
    if st.button("Analizar Sentimiento"):
        if frase:
            frase_en = traducir_to_eng(frase)
            sentimiento = ola.polarity_scores(frase_en)
            st.write("Frase original:", frase)
            st.write("Frase traducida al inglés:", frase_en)
            st.write("Puntuación de sentimiento:", sentimiento)
        else:
            st.write("Por favor, introduce una frase.")

def chunk():
    nltk.download('averaged_perceptron_tagger')
    nltk.download('punkt')

    translator = Translator()
    text_spanish = st.text_input("Introduce una frase para ser analizada")
    if st.button("Analizar Chunking"):
        if text_spanish:
            text_english = translator.translate(text_spanish, src='es', dest='en').text
            words = word_tokenize(text_english)
            tagged = nltk.pos_tag(words)
            grammar = "NP: {<DT>?<JJ>*<NN>}"
            parser = RegexpParser(grammar)
            result = parser.parse(tagged)
            st.write("Frase traducida y analizada:", result)
            st.write("Ver el árbol de chunks (requiere entorno gráfico).")
        else:
            st.write("Por favor, introduce una frase.")

def reconocer_entidades():
    nltk.download('punkt')
    nltk.download('averaged_perceptron_tagger')
    nltk.download('maxent_ne_chunker')
    nltk.download('words')

    def extract_text_from_docx(file): 
        doc = docx.Document(file)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        return '\n'.join(text)

    def get_named_entities(text): 
        tokens = word_tokenize(text)
        pos_tags = pos_tag(tokens)
        ner_chunks = ne_chunk(pos_tags)
        
        named_entities = []
        for chunk in ner_chunks:
            if isinstance(chunk, Tree):
                entity = " ".join([token for token, pos in chunk.leaves()])
                entity_type = chunk.label()
                named_entities.append((entity, entity_type))
        
        return named_entities

    st.title("Reconocimiento de Entidades Nombradas con NLTK")
    
    st.subheader("Cargar archivo DOCX")
    docx_file = st.file_uploader("Sube tu archivo DOCX", type=["docx"])

    if docx_file is not None:
        raw_text = extract_text_from_docx(docx_file)
        st.text_area("Texto del archivo DOCX", raw_text, height=200)
        
        if st.button("Reconocer Entidades"):
            named_entities = get_named_entities(raw_text)
            st.subheader("Entidades Reconocidas")
            if named_entities:
                for entity, entity_type in named_entities:
                    st.write(f'{entity} ({entity_type})')
            else:
                st.write("No se encontraron entidades nombradas.")


def convertir_py_a_exe(archivo_py):
    try:
        subprocess.run(['pyinstaller', '--onefile', archivo_py])
        st.write("¡El archivo .exe ha sido creado exitosamente!")
    except Exception as e:
        st.write(f"Se produjo un error al intentar convertir el archivo a .exe: {e}")

def archivo_py_a_exe():
    archivo_py = st.file_uploader("Sube tu archivo .py para convertir a .exe", type=["py"])
    if st.button("Convertir a .exe"):
        if archivo_py:
            convertir_py_a_exe(archivo_py.name)
        else:
            st.write("Por favor, sube un archivo .py.")

################################################################
# Menú principal de Streamlit
st.title("Menú de actividades: Procesamiento de lenguaje natural")

menu = ["Audio a Texto", "Texto a Audio", "Archivo a Audio", "ChatGPT", "Análisis de Sentimientos", "Chunking", "Reconocimiento de Entidades", "Convertir .py a .exe"]
choice = st.sidebar.selectbox("Selecciona una función", menu)

if choice == "Audio a Texto":
    grabar_audio_texto()
elif choice == "Texto a Audio":
    texto_audio()
elif choice == "Archivo a Audio":
    audio_texto_audio()
elif choice == "ChatGPT":
    chatGPT()
elif choice == "Análisis de Sentimientos":
    analizador_sentiment()
elif choice == "Chunking":
    chunk()
elif choice == "Reconocimiento de Entidades":
    reconocer_entidades()
elif choice == "Convertir .py a .exe":
    archivo_py_a_exe()

################################################################

